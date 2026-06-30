import os
import re
import shutil
import uuid

import pypdfium2 as pdfium
import pytesseract

from docx import Document as WordDocument
from PIL import (
    Image,
    ImageOps,
    UnidentifiedImageError
)
from pypdf import PdfReader
from pytesseract import (
    TesseractError,
    TesseractNotFoundError
)
from werkzeug.utils import secure_filename


IMAGE_EXTENSIONS = {
    "png",
    "jpg",
    "jpeg",
    "webp"
}


class DocumentProcessingError(Exception):
    status_code = 422


class OCRUnavailableError(
    DocumentProcessingError
):
    status_code = 503


class DocumentLimitError(
    DocumentProcessingError
):
    status_code = 422


class UnreadableDocumentError(
    DocumentProcessingError
):
    status_code = 422


def get_file_extension(filename):
    if not filename or "." not in filename:
        return ""

    return filename.rsplit(
        ".",
        1
    )[1].lower()


def is_allowed_file(
    filename,
    allowed_extensions
):
    extension = get_file_extension(
        filename
    )

    return extension in allowed_extensions


def save_uploaded_file(
    file_storage,
    upload_folder
):
    os.makedirs(
        upload_folder,
        exist_ok=True
    )

    original_filename = (
        file_storage.filename
        or "document"
    )

    safe_filename = secure_filename(
        original_filename
    )

    extension = get_file_extension(
        safe_filename
    )

    if not extension:
        raise DocumentProcessingError(
            "The uploaded file has no supported extension."
        )

    stored_filename = (
        f"{uuid.uuid4().hex}.{extension}"
    )

    file_path = os.path.join(
        upload_folder,
        stored_filename
    )

    file_storage.save(file_path)

    return {
        "original_filename": original_filename,
        "stored_filename": stored_filename,
        "extension": extension,
        "file_path": file_path,
        "file_size": os.path.getsize(
            file_path
        )
    }


def extract_text_from_file(
    file_path,
    extension,
    ocr_settings=None
):
    extension = extension.lower()

    settings = normalize_ocr_settings(
        ocr_settings
    )

    if extension == "txt":
        raw_text = extract_text_from_txt(
            file_path
        )

        result = {
            "text": raw_text,
            "method": "native",
            "ocr_used": False,
            "pages_processed": 1,
            "native_pages": 1,
            "ocr_pages": 0
        }

    elif extension == "docx":
        raw_text = extract_text_from_docx(
            file_path
        )

        result = {
            "text": raw_text,
            "method": "native",
            "ocr_used": False,
            "pages_processed": 1,
            "native_pages": 1,
            "ocr_pages": 0
        }

    elif extension == "pdf":
        result = extract_text_from_pdf(
            file_path,
            settings
        )

    elif extension in IMAGE_EXTENSIONS:
        result = extract_text_from_image(
            file_path,
            settings
        )

    else:
        raise DocumentProcessingError(
            f"Unsupported file type: {extension}"
        )

    cleaned_text = clean_extracted_text(
        result.get("text", "")
    )

    if not cleaned_text:
        raise UnreadableDocumentError(
            "No readable text was found. "
            "Try a clearer image, a higher-resolution scan, "
            "or another document."
        )

    result["text"] = cleaned_text
    result["text_length"] = len(
        cleaned_text
    )

    return result


def extract_text_from_txt(file_path):
    try:
        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as file:
            return file.read()

    except UnicodeDecodeError:
        with open(
            file_path,
            "r",
            encoding="latin-1"
        ) as file:
            return file.read()


def extract_text_from_pdf(
    file_path,
    settings
):
    try:
        reader = PdfReader(
            file_path
        )
    except Exception as error:
        raise UnreadableDocumentError(
            f"The PDF could not be opened: {error}"
        ) from error

    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception:
            unlocked = 0

        if not unlocked:
            raise UnreadableDocumentError(
                "Password-protected PDFs are not supported."
            )

    page_count = len(
        reader.pages
    )

    if page_count == 0:
        raise UnreadableDocumentError(
            "The PDF contains no pages."
        )

    if page_count > settings["max_pdf_pages"]:
        raise DocumentLimitError(
            "The PDF has "
            f"{page_count} pages. "
            "The current safety limit is "
            f"{settings['max_pdf_pages']} pages."
        )

    native_pages = {}
    pages_for_ocr = []

    for page_index, page in enumerate(
        reader.pages
    ):
        page_text = ""

        try:
            page_text = (
                page.extract_text()
                or ""
            )
        except Exception:
            page_text = ""

        cleaned_page_text = (
            clean_extracted_text(
                page_text
            )
        )

        character_count = count_text_characters(
            cleaned_page_text
        )

        has_images = pdf_page_has_images(
            page
        )

        if (
            character_count
            >= settings[
                "min_text_characters"
            ]
        ):
            native_pages[
                page_index
            ] = cleaned_page_text

        elif (
            settings["ocr_enabled"]
            and has_images
        ):
            pages_for_ocr.append(
                page_index
            )

        elif cleaned_page_text:
            native_pages[
                page_index
            ] = cleaned_page_text

    if (
        pages_for_ocr
        and len(pages_for_ocr)
        > settings["ocr_max_pdf_pages"]
    ):
        raise DocumentLimitError(
            "This PDF needs OCR on "
            f"{len(pages_for_ocr)} pages, "
            "but the current OCR safety limit is "
            f"{settings['ocr_max_pdf_pages']} pages."
        )

    ocr_pages = {}

    if pages_for_ocr:
        ensure_ocr_available(
            settings
        )

        ocr_pages = run_pdf_ocr(
            file_path=file_path,
            page_indexes=pages_for_ocr,
            settings=settings
        )

    combined_pages = []

    for page_index in range(
        page_count
    ):
        page_text = (
            ocr_pages.get(page_index)
            or native_pages.get(page_index)
            or ""
        )

        page_text = clean_extracted_text(
            page_text
        )

        if page_text:
            combined_pages.append(
                f"[Page {page_index + 1}]\n"
                f"{page_text}"
            )

    if not combined_pages:
        if not settings["ocr_enabled"]:
            raise UnreadableDocumentError(
                "No readable text was found in this PDF. "
                "OCR is currently disabled."
            )

        if not pages_for_ocr:
            raise UnreadableDocumentError(
                "No readable text or OCR-compatible page images "
                "were found in this PDF."
            )

        raise UnreadableDocumentError(
            "OCR could not detect readable text in this PDF."
        )

    native_page_count = len(
        native_pages
    )

    ocr_page_count = len(
        [
            text
            for text in ocr_pages.values()
            if clean_extracted_text(text)
        ]
    )

    if (
        native_page_count > 0
        and ocr_page_count > 0
    ):
        method = "hybrid"

    elif ocr_page_count > 0:
        method = "ocr"

    else:
        method = "native"

    return {
        "text": "\n\n".join(
            combined_pages
        ),
        "method": method,
        "ocr_used": (
            ocr_page_count > 0
        ),
        "pages_processed": page_count,
        "native_pages": native_page_count,
        "ocr_pages": ocr_page_count
    }


def extract_text_from_docx(file_path):
    try:
        document = WordDocument(
            file_path
        )
    except Exception as error:
        raise UnreadableDocumentError(
            f"The DOCX file could not be opened: {error}"
        ) from error

    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if cells:
                paragraphs.append(
                    " | ".join(cells)
                )

    return "\n\n".join(
        paragraphs
    )


def extract_text_from_image(
    file_path,
    settings
):
    if not settings["ocr_enabled"]:
        raise OCRUnavailableError(
            "Image uploads require OCR, but OCR is disabled."
        )

    ensure_ocr_available(
        settings
    )

    image = open_image_safely(
        file_path,
        settings["max_image_pixels"]
    )

    try:
        prepared_image = prepare_image_for_ocr(
            image,
            settings["max_image_pixels"]
        )

        try:
            text = run_tesseract(
                prepared_image,
                settings
            )
        finally:
            if prepared_image is not image:
                prepared_image.close()

    finally:
        image.close()

    text = clean_extracted_text(
        text
    )

    if not text:
        raise UnreadableDocumentError(
            "OCR could not detect readable text in this image. "
            "Try a sharper image with stronger contrast."
        )

    return {
        "text": text,
        "method": "ocr",
        "ocr_used": True,
        "pages_processed": 1,
        "native_pages": 0,
        "ocr_pages": 1
    }


def run_pdf_ocr(
    file_path,
    page_indexes,
    settings
):
    try:
        pdf = pdfium.PdfDocument(
            file_path
        )
    except Exception as error:
        raise UnreadableDocumentError(
            f"The PDF could not be rendered for OCR: {error}"
        ) from error

    scale = (
        settings["pdf_dpi"]
        / 72.0
    )

    results = {}

    try:
        for page_index in page_indexes:
            page = None
            bitmap = None
            rendered_image = None
            image = None
            prepared_image = None

            try:
                page = pdf[
                    page_index
                ]

                bitmap = page.render(
                    scale=scale
                )

                rendered_image = (
                    bitmap.to_pil()
                )

                image = rendered_image.copy()

                enforce_image_pixel_limit(
                    image,
                    settings[
                        "max_image_pixels"
                    ]
                )

                prepared_image = (
                    prepare_image_for_ocr(
                        image,
                        settings[
                            "max_image_pixels"
                        ]
                    )
                )

                results[
                    page_index
                ] = clean_extracted_text(
                    run_tesseract(
                        prepared_image,
                        settings
                    )
                )

            except DocumentProcessingError:
                raise

            except Exception as error:
                raise UnreadableDocumentError(
                    "OCR failed while reading "
                    f"PDF page {page_index + 1}: {error}"
                ) from error

            finally:
                if (
                    prepared_image is not None
                    and prepared_image is not image
                ):
                    prepared_image.close()

                if image is not None:
                    image.close()

                if rendered_image is not None:
                    rendered_image.close()

                if bitmap is not None:
                    try:
                        bitmap.close()
                    except Exception:
                        pass

                if page is not None:
                    try:
                        page.close()
                    except Exception:
                        pass

    finally:
        try:
            pdf.close()
        except Exception:
            pass

    return results


def run_tesseract(
    image,
    settings
):
    command = ensure_ocr_available(
        settings
    )

    pytesseract.pytesseract.tesseract_cmd = (
        command
    )

    try:
        return pytesseract.image_to_string(
            image,
            lang=settings["language"],
            config=settings[
                "tesseract_config"
            ]
        )

    except TesseractNotFoundError as error:
        raise OCRUnavailableError(
            "Tesseract OCR was not found. "
            "Install Tesseract or configure TESSERACT_CMD."
        ) from error

    except TesseractError as error:
        raise UnreadableDocumentError(
            "Tesseract OCR could not process the file. "
            f"Details: {error}"
        ) from error


def normalize_ocr_settings(
    settings
):
    raw = settings or {}

    return {
        "ocr_enabled": bool(
            raw.get(
                "OCR_ENABLED",
                True
            )
        ),
        "tesseract_cmd": (
            raw.get(
                "TESSERACT_CMD"
            )
            or None
        ),
        "language": (
            str(
                raw.get(
                    "OCR_LANGUAGE",
                    "eng"
                )
            ).strip()
            or "eng"
        ),
        "tesseract_config": (
            str(
                raw.get(
                    "OCR_TESSERACT_CONFIG",
                    "--oem 3 --psm 3"
                )
            ).strip()
        ),
        "min_text_characters": int(
            raw.get(
                "OCR_MIN_TEXT_CHARACTERS",
                40
            )
        ),
        "max_pdf_pages": int(
            raw.get(
                "MAX_PDF_PAGES",
                100
            )
        ),
        "ocr_max_pdf_pages": int(
            raw.get(
                "OCR_MAX_PDF_PAGES",
                25
            )
        ),
        "pdf_dpi": int(
            raw.get(
                "OCR_PDF_DPI",
                200
            )
        ),
        "max_image_pixels": int(
            raw.get(
                "OCR_MAX_IMAGE_PIXELS",
                25_000_000
            )
        )
    }


def get_ocr_status(settings=None):
    normalized = normalize_ocr_settings(
        settings
    )

    command = resolve_tesseract_command(
        normalized["tesseract_cmd"]
    )

    return {
        "enabled": normalized[
            "ocr_enabled"
        ],
        "available": bool(command),
        "language": normalized[
            "language"
        ]
    }


def ensure_ocr_available(settings):
    if not settings["ocr_enabled"]:
        raise OCRUnavailableError(
            "OCR is disabled. Set OCR_ENABLED=true "
            "to process scanned documents and images."
        )

    command = resolve_tesseract_command(
        settings.get(
            "tesseract_cmd"
        )
    )

    if not command:
        raise OCRUnavailableError(
            "Tesseract OCR was not found. "
            "Install Tesseract, restart PyCharm, "
            "or set TESSERACT_CMD in the .env file."
        )

    return command


def resolve_tesseract_command(
    configured_command=None
):
    candidates = []

    if configured_command:
        candidates.append(
            os.path.expandvars(
                os.path.expanduser(
                    configured_command
                )
            )
        )

    discovered = shutil.which(
        "tesseract"
    )

    if discovered:
        candidates.append(
            discovered
        )

    local_app_data = os.getenv(
        "LOCALAPPDATA",
        ""
    )

    program_files = os.getenv(
        "ProgramFiles",
        r"C:\Program Files"
    )

    program_files_x86 = os.getenv(
        "ProgramFiles(x86)",
        r"C:\Program Files (x86)"
    )

    candidates.extend([
        os.path.join(
            program_files,
            "Tesseract-OCR",
            "tesseract.exe"
        ),
        os.path.join(
            program_files_x86,
            "Tesseract-OCR",
            "tesseract.exe"
        ),
        os.path.join(
            local_app_data,
            "Programs",
            "Tesseract-OCR",
            "tesseract.exe"
        )
    ])

    for candidate in candidates:
        if (
            candidate
            and os.path.isfile(
                candidate
            )
        ):
            return candidate

    return None


def open_image_safely(
    file_path,
    max_pixels
):
    try:
        image = Image.open(
            file_path
        )

        enforce_image_pixel_limit(
            image,
            max_pixels
        )

        image.load()

        return image

    except (
        UnidentifiedImageError,
        OSError
    ) as error:
        raise UnreadableDocumentError(
            f"The image could not be opened: {error}"
        ) from error


def enforce_image_pixel_limit(
    image,
    max_pixels
):
    width, height = image.size

    if width <= 0 or height <= 0:
        raise UnreadableDocumentError(
            "The image has invalid dimensions."
        )

    pixel_count = (
        width * height
    )

    if pixel_count > max_pixels:
        raise DocumentLimitError(
            "The image contains "
            f"{pixel_count:,} pixels. "
            "The current safety limit is "
            f"{max_pixels:,} pixels."
        )


def prepare_image_for_ocr(
    image,
    max_pixels
):
    prepared = ImageOps.exif_transpose(
        image
    ).convert(
        "L"
    )

    prepared = ImageOps.autocontrast(
        prepared
    )

    width, height = prepared.size
    pixel_count = width * height

    if (
        max(width, height) < 1400
        and pixel_count * 4 <= max_pixels
    ):
        resized = prepared.resize(
            (
                width * 2,
                height * 2
            ),
            Image.Resampling.LANCZOS
        )

        prepared.close()
        prepared = resized

    return prepared


def pdf_page_has_images(page):
    try:
        return bool(
            list(page.images)
        )
    except Exception:
        pass

    try:
        resources = page.get(
            "/Resources"
        )

        if not resources:
            return False

        resources = resources.get_object()

        x_objects = resources.get(
            "/XObject"
        )

        if not x_objects:
            return False

        x_objects = x_objects.get_object()

        for item in x_objects.values():
            try:
                obj = item.get_object()

                if obj.get(
                    "/Subtype"
                ) == "/Image":
                    return True
            except Exception:
                continue

    except Exception:
        return False

    return False


def count_text_characters(text):
    return sum(
        1
        for character in text
        if character.isalnum()
    )


def clean_extracted_text(text):
    if not text:
        return ""

    text = text.replace(
        "\x00",
        " "
    )

    text = text.replace(
        "\r\n",
        "\n"
    )

    text = text.replace(
        "\r",
        "\n"
    )

    lines = []

    for line in text.split("\n"):
        cleaned_line = re.sub(
            r"[ \t]+",
            " ",
            line
        ).strip()

        lines.append(
            cleaned_line
        )

    cleaned_text = "\n".join(
        lines
    )

    cleaned_text = re.sub(
        r"\n{3,}",
        "\n\n",
        cleaned_text
    )

    return cleaned_text.strip()


def split_text_into_chunks(
    text,
    chunk_size=450,
    overlap=80
):
    if chunk_size <= 0:
        raise ValueError(
            "Chunk size must be greater than zero."
        )

    if overlap < 0:
        raise ValueError(
            "Chunk overlap cannot be negative."
        )

    if overlap >= chunk_size:
        raise ValueError(
            "Chunk overlap must be smaller "
            "than chunk size."
        )

    words = text.split()

    if not words:
        return []

    chunks = []
    start = 0

    while start < len(words):
        end = min(
            start + chunk_size,
            len(words)
        )

        chunk = " ".join(
            words[start:end]
        ).strip()

        if chunk:
            chunks.append(
                chunk
            )

        if end >= len(words):
            break

        start = end - overlap

    return chunks
